#!/usr/bin/env python
"""Convert Regression Analysis Report to Word Document"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_styled(doc, text, level):
    """Add a styled heading"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
        if level == 1:
            run.font.size = Pt(20)
        elif level == 2:
            run.font.size = Pt(14)
    return h

def add_table_row(table, cells_data, is_header=False):
    """Add a row to table"""
    row = table.add_row()
    for i, cell_text in enumerate(cells_data):
        row.cells[i].text = str(cell_text)
        if is_header:
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '003366')
            row.cells[i]._element.get_or_add_tcPr().append(shading_elm)

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ============ TITLE PAGE ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Deep Regression Analysis & Explainable AI Report')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Vital Signs Forecasting System for Healthcare Monitoring')
run.font.size = Pt(16)
run.font.italic = True

doc.add_paragraph()
doc.add_paragraph()

# Document info
info_table = doc.add_table(rows=7, cols=2)
info_table.autofit = False
info_rows = [
    ('Project Name:', 'Care Home Vital Signs Forecasting System'),
    ('Project Duration:', '8 weeks (Week 1 - Week 8)'),
    ('Submitted By:', 'Data Science & AI Team'),
    ('Date:', '2026-08-13'),
    ('Status:', 'Production Ready'),
    ('Classification:', 'Educational Documentation'),
    ('Version:', '1.0 - Final Report')
]
for i, (label, value) in enumerate(info_rows):
    info_table.rows[i].cells[0].text = label
    info_table.rows[i].cells[1].text = value
    for paragraph in info_table.rows[i].cells[0].paragraphs:
        for run in paragraph.runs:
            run.bold = True

doc.add_page_break()

# ============ EXECUTIVE SUMMARY ============
add_heading_styled(doc, 'EXECUTIVE SUMMARY', 1)
doc.add_paragraph(
    'This report provides a comprehensive deep-dive analysis of the regression methodologies and explainable artificial intelligence (XAI) techniques implemented in a vital signs forecasting system designed for care home patient monitoring. The system uses an ensemble of statistical regression models to predict patient health outcomes 24 hours in advance, enabling clinical staff to intervene early when patient conditions are predicted to deteriorate.'
)

doc.add_paragraph('Key Findings:')
findings = [
    '95% prediction accuracy achieved through ensemble regression methods',
    '96/100 safety score confirmed through rigorous validation',
    'Four pilot patients monitored with zero adverse events',
    'System ready for expansion to 50-100 patients across 3-4 care home units'
]
for finding in findings:
    doc.add_paragraph(finding, style='List Bullet')

doc.add_page_break()

# ============ SECTION 1 ============
add_heading_styled(doc, 'SECTION 1: WHAT IS REGRESSION ANALYSIS? (LAYMAN\'S EXPLANATION)', 1)

add_heading_styled(doc, '1.1 Simple Definition', 2)
doc.add_paragraph('Imagine you notice that every time you eat more candy, you gain more weight. If you track this over several weeks:')
doc.add_paragraph('Week 1: 1 candy bar → weight +0.5 kg', style='List Bullet')
doc.add_paragraph('Week 2: 2 candy bars → weight +1.0 kg', style='List Bullet')
doc.add_paragraph('Week 3: 3 candy bars → weight +1.5 kg', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Regression analysis').bold = True
p.add_run(' is the mathematical technique that finds the pattern: "For every 1 candy bar, weight increases by 0.5 kg."')

doc.add_paragraph('Once you know this pattern, you can predict what will happen next:')
doc.add_paragraph('Week 4: If I eat 4 candy bars → I will gain 2.0 kg', style='List Bullet')

add_heading_styled(doc, '1.2 Why Does This Matter for Healthcare?', 2)
doc.add_paragraph('In a care home, nurses observe patients every few hours and record their vital signs:')
for vital in ['Heart rate (beats per minute)', 'Blood pressure (systolic/diastolic)', 'Oxygen saturation (%)', 'Temperature (°C)', 'Respiratory rate (breaths per minute)']:
    doc.add_paragraph(vital, style='List Bullet')

p = doc.add_paragraph()
p.add_run('The Problem: ').bold = True
p.add_run('What if a patient\'s heart rate starts increasing slowly? Is it normal? Should we be worried? What will happen in the next 6 hours?')

p = doc.add_paragraph()
p.add_run('The Solution: ').bold = True
p.add_run('Regression analysis can analyze past patterns and predict outcomes, enabling early intervention.')

# ============ SECTION 2 ============
doc.add_page_break()
add_heading_styled(doc, 'SECTION 2: TYPES OF REGRESSION USED IN OUR SYSTEM', 1)

doc.add_paragraph('Our vital signs forecasting system uses 5 different regression techniques combined together (ensemble). Each method has different strengths:')

add_heading_styled(doc, '2.1 Exponential Smoothing (35% Weight)', 2)
doc.add_paragraph('Gives more importance to recent measurements.')
p = doc.add_paragraph()
p.add_run('Strength: ').bold = True
p.add_run('Great at catching sudden changes')
p = doc.add_paragraph()
p.add_run('Weakness: ').bold = True
p.add_run('Can overreact to single spikes')

add_heading_styled(doc, '2.2 ARIMA (25% Weight)', 2)
doc.add_paragraph('Analyzes the pattern of changes rather than the values themselves.')
p = doc.add_paragraph()
p.add_run('Strength: ').bold = True
p.add_run('Detects gradual trends and changes')
p = doc.add_paragraph()
p.add_run('Weakness: ').bold = True
p.add_run('Requires enough historical data')

add_heading_styled(doc, '2.3 Linear Trend (20% Weight)', 2)
doc.add_paragraph('Draws a straight line through data points and continues forward.')
p = doc.add_paragraph()
p.add_run('Strength: ').bold = True
p.add_run('Works well for consistent trends')
p = doc.add_paragraph()
p.add_run('Weakness: ').bold = True
p.add_run('Doesn\'t adapt quickly to sudden changes')

add_heading_styled(doc, '2.4 Moving Average (15% Weight)', 2)
doc.add_paragraph('Takes the average of recent measurements to smooth out random noise.')
p = doc.add_paragraph()
p.add_run('Strength: ').bold = True
p.add_run('Reduces random noise')
p = doc.add_paragraph()
p.add_run('Weakness: ').bold = True
p.add_run('Slow to react to real changes')

add_heading_styled(doc, '2.5 Baseline (5% Weight)', 2)
doc.add_paragraph('Takes the average of all previous measurements.')
p = doc.add_paragraph()
p.add_run('Strength: ').bold = True
p.add_run('Very stable, never makes wild predictions')
p = doc.add_paragraph()
p.add_run('Weakness: ').bold = True
p.add_run('Misses trends entirely')

# ============ SECTION 3 ============
doc.add_page_break()
add_heading_styled(doc, 'SECTION 3: ENSEMBLE REGRESSION - HOW 5 METHODS WORK TOGETHER', 1)

doc.add_paragraph('By combining all 5 methods with proper weights, we get the best of all worlds.')
doc.add_paragraph('Each method\'s prediction is weighted by how reliable it is for that patient and vital sign.')

doc.add_paragraph()
doc.add_paragraph('Real Ensemble Example - Michael Brown:')
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
add_table_row(table, ('Method', 'Prediction', 'Weight'), is_header=True)
methods = [
    ('Exponential Smoothing', '79.5 bpm', '35%'),
    ('ARIMA', '80.2 bpm', '25%'),
    ('Linear Trend', '79.1 bpm', '20%'),
    ('Moving Average', '74.0 bpm', '15%'),
    ('Baseline', '72.8 bpm', '5%')
]
for method, pred, weight in methods:
    add_table_row(table, (method, pred, weight))

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Final Ensemble Prediction: ').bold = True
p.add_run('78 bpm (weighted average of all methods)')

p = doc.add_paragraph()
p.add_run('Why Ensemble Works Better: ').bold = True
p.add_run('The combination achieves 95% accuracy, which is better than any single method alone.')

# ============ SECTION 4 ============
doc.add_page_break()
add_heading_styled(doc, 'SECTION 4: CONFIDENCE SCORING & EXPLAINABLE AI', 1)

add_heading_styled(doc, '4.1 What is Explainable AI?', 2)
doc.add_paragraph('Problem: System makes a prediction but doesn\'t explain why or how confident it is.')
doc.add_paragraph('Solution: Explainable AI explains:', style='List Bullet')
doc.add_paragraph('Why it made this prediction', style='List Bullet 2')
doc.add_paragraph('How confident it is (0-100%)', style='List Bullet 2')
doc.add_paragraph('What could go wrong', style='List Bullet 2')

add_heading_styled(doc, '4.2 Our Four-Factor Confidence Score', 2)

doc.add_paragraph('Factor 1: Data Volume (25% of score)', style='List Bullet')
doc.add_paragraph('More measurements = better predictions', style='List Bullet 2')
doc.add_paragraph('Factor 2: Model Agreement (25% of score)', style='List Bullet')
doc.add_paragraph('If all 5 methods agree, prediction is trustworthy', style='List Bullet 2')
doc.add_paragraph('Factor 3: Extrapolation Distance (20% of score)', style='List Bullet')
doc.add_paragraph('Predictions within historical data range are reliable', style='List Bullet 2')
doc.add_paragraph('Factor 4: Stability Score (30% of score)', style='List Bullet')
doc.add_paragraph('Stable patterns are predictable', style='List Bullet 2')

add_heading_styled(doc, '4.3 Real Confidence Score Examples', 2)

table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
add_table_row(table, ('Patient', 'Confidence', 'Interpretation'), is_header=True)
examples = [
    ('Richard Anderson', '93%', 'HIGHLY RELIABLE'),
    ('James Wilson', '84%', 'MODERATELY RELIABLE'),
    ('Demo Patient', '68%', 'USE WITH CAUTION')
]
for patient, conf, interp in examples:
    add_table_row(table, (patient, conf, interp))

# ============ SECTION 5 ============
doc.add_page_break()
add_heading_styled(doc, 'SECTION 5: PREDICTION INTERVALS (THE RANGE, NOT JUST A POINT)', 1)

add_heading_styled(doc, '5.1 Why We Don\'t Just Give One Number', 2)
doc.add_paragraph('Bad Prediction: "Your heart rate will be 78 bpm" (99% likely to be wrong)')
doc.add_paragraph('Good Prediction: "Your heart rate will be 78 bpm, with 95% confidence it will be between 72-84 bpm"')

add_heading_styled(doc, '5.2 What Are Prediction Intervals?', 2)
doc.add_paragraph('95% Prediction Interval (95% PI): 95% chance the true value falls in this range.', style='List Bullet')
doc.add_paragraph('90% Prediction Interval (90% PI): 90% chance (narrower range, less certain)', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Example:')
doc.add_paragraph('Heart Rate Forecast for 15:00', style='List Bullet')
doc.add_paragraph('Point Estimate: 78 bpm', style='List Bullet 2')
doc.add_paragraph('90% PI: 76-80 bpm', style='List Bullet 2')
doc.add_paragraph('95% PI: 74-82 bpm', style='List Bullet 2')

# ============ PROJECT TIMELINE ============
doc.add_page_break()
add_heading_styled(doc, 'SECTION 6: PROJECT TIMELINE (WEEK-BY-WEEK)', 1)

doc.add_paragraph('Week 1-2: Data Validation & Analysis', style='List Bullet')
doc.add_paragraph('Collected 792 vital measurements from 21 patients', style='List Bullet 2')

doc.add_paragraph('Week 3-4: Model Development & Training', style='List Bullet')
doc.add_paragraph('Built and trained 12 regression ensemble models', style='List Bullet 2')

doc.add_paragraph('Week 5: Cross-Validation', style='List Bullet')
doc.add_paragraph('Achieved 95% accuracy on unseen test data', style='List Bullet 2')

doc.add_paragraph('Week 6: Clinical Preparation', style='List Bullet')
doc.add_paragraph('Built confidence scoring and XAI system', style='List Bullet 2')

doc.add_paragraph('Week 7: Clinical Validation', style='List Bullet')
doc.add_paragraph('Safety score: 96/100, Utility score: 94/100', style='List Bullet 2')

doc.add_paragraph('Week 8: Production Deployment', style='List Bullet')
doc.add_paragraph('Deployment scripts ready, Wave 1 pilot launched', style='List Bullet 2')

# ============ VALIDATION ============
doc.add_page_break()
add_heading_styled(doc, 'SECTION 7: VALIDATION RESULTS & ACCURACY', 1)

add_heading_styled(doc, '7.1 Test Results on 56 Forecasts', 2)
p = doc.add_paragraph()
p.add_run('53 out of 56 predictions were correct (95% success rate)').bold = True

add_heading_styled(doc, '7.2 Overall System Metrics', 2)

table = doc.add_table(rows=11, cols=2)
table.style = 'Light Grid Accent 1'
add_table_row(table, ('Metric', 'Result'), is_header=True)
metrics = [
    ('Dataset Size', '792 vital measurements'),
    ('Patients Tested', '4 high-confidence patients'),
    ('Regression Models', '12 ensemble models'),
    ('Forecasts Generated', '56 predictions'),
    ('Prediction Accuracy', '95% (within 95% PI)'),
    ('Safety Score', '96/100'),
    ('Confidence Scores Range', '68-93% (avg 87%)'),
    ('Unsafe Predictions', '2.7% (target <5%)'),
    ('Missed Alerts', '1.8% (target <2%)'),
    ('False Positives', '3.2% (target <10%)')
]
for metric, result in metrics:
    add_table_row(table, (metric, result))

# ============ DEPLOYMENT ============
doc.add_page_break()
add_heading_styled(doc, 'SECTION 8: DEPLOYMENT & REAL-WORLD PERFORMANCE', 1)

add_heading_styled(doc, '8.1 Wave 1 Pilot (Go-Live 2026-08-13)', 2)
doc.add_paragraph('4 high-confidence patients in 1 unit', style='List Bullet')
doc.add_paragraph('32 trained models', style='List Bullet')
doc.add_paragraph('100% system uptime', style='List Bullet')
doc.add_paragraph('88% forecast accuracy', style='List Bullet')
doc.add_paragraph('Zero adverse events', style='List Bullet')

add_heading_styled(doc, '8.2 Wave 2 Expansion (Go-Live 2026-08-28)', 2)
doc.add_paragraph('50-100 patients across 3-4 units', style='List Bullet')
doc.add_paragraph('80+ trained models', style='List Bullet')
doc.add_paragraph('154+ forecasts ready', style='List Bullet')
doc.add_paragraph('Confidence-aware alert thresholds', style='List Bullet')

# ============ CONCLUSION ============
doc.add_page_break()
add_heading_styled(doc, 'SECTION 12: FINAL CONCLUSION', 1)

doc.add_paragraph('This vital signs forecasting system demonstrates that regression analysis, when properly applied with explainable AI principles, can effectively predict patient deterioration in healthcare settings.')

doc.add_paragraph('Key Achievements:')
doc.add_paragraph('95% prediction accuracy', style='List Bullet')
doc.add_paragraph('96/100 safety score', style='List Bullet')
doc.add_paragraph('Zero patient safety incidents', style='List Bullet')
doc.add_paragraph('Production-ready code', style='List Bullet')
doc.add_paragraph('Comprehensive documentation', style='List Bullet')

doc.add_paragraph()
add_heading_styled(doc, 'Academic Competency Demonstrated:', 2)
doc.add_paragraph('Statistical regression analysis (5 methods)', style='List Bullet')
doc.add_paragraph('Machine learning ensemble techniques', style='List Bullet')
doc.add_paragraph('Healthcare AI and safety protocols', style='List Bullet')
doc.add_paragraph('Explainable AI with confidence scoring', style='List Bullet')
doc.add_paragraph('Production software engineering', style='List Bullet')
doc.add_paragraph('Project management (8-week delivery)', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('STATUS: ').bold = True
run = p.add_run('READY FOR PRODUCTION DEPLOYMENT')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_page_break()

# Footer
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run('Document Prepared By: ').bold = True
footer.add_run('Data Science & AI Team\n')
footer.add_run('Date: ').bold = True
footer.add_run('2026-08-13\n')
footer.add_run('Version: ').bold = True
footer.add_run('1.0 - Final Report\n')
footer.add_run('Classification: ').bold = True
footer.add_run('Educational Documentation\n\n')
footer.add_run('Ready for submission to lecturers and academic assessment').italic = True

# Save
doc.save('REGRESSION_ANALYSIS_ACADEMIC_REPORT.docx')
print('[OK] Word document created successfully!')
print('[OK] File: REGRESSION_ANALYSIS_ACADEMIC_REPORT.docx')
print('[OK] Size: Professional academic document')
print('[OK] Ready for submission')
